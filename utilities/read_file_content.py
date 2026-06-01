from collections import Counter
from typing import List
from io import BytesIO
import logging

from docx import Document as DocxDocument
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from routers.bloom_api import filter_missing_words

logger = logging.getLogger(__name__)

# This is the string containing all unwanted characters
special_characters = r"""೧^l=F–೬B#yJwfz•+2umE<'!CxULvr]8o೦VNd0hH'_>)- :sYQ7.g9n%W,G`1…"&?6೯I"೮೨Tb"@೭೫ʼKX4೪[iDScM;*t'{5k/pa(PAeZ~O3R|j}q೩$"""

# Create a translation table to remove unwanted characters
translation_table = str.maketrans('', '', special_characters)


async def extract_words_from_txt(file: UploadFile) -> List[str]:
    """
    Extracts words from a .txt file.

    Args:
        file (UploadFile): The uploaded .txt file.

    Returns:
        List[str]: A list of words extracted from the file.
    """
    try:
        if not file.filename.endswith('.txt'):
            logger.warning(f"Invalid TXT file extension: {file.filename}")
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .txt file.")

        contents = await file.read()
        
        # Try UTF-8 first, then fallback to other encodings
        try:
            text = contents.decode('utf-8')
        except UnicodeDecodeError:
            # Try other common encodings
            for encoding in ['utf-16', 'iso-8859-1', 'cp1252']:
                try:
                    text = contents.decode(encoding)
                    logger.info(f"Successfully decoded file with {encoding}")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise HTTPException(status_code=400, detail="Could not decode file. Please ensure it's a valid text file with UTF-8 encoding.")
        
        words = text.split()
        logger.info(f"Successfully extracted {len(words)} words from TXT file: {file.filename}")
        return words
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reading TXT file {file.filename}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"Error processing TXT file: {str(e)}")


def clean_words(words: List[str]) -> List[str]:
    """
    Cleans the list of words by removing any characters found in the string `p`.

    Args:
        words (List[str]): The list of words extracted from the file.

    Returns:
        List[str]: The cleaned list of words.
    """
    cleaned_words = []
    for word in words:
        # Remove any characters found in the string `p`
        cleaned_word = word.translate(translation_table)
        # Add the cleaned word to the list if it's not empty
        if cleaned_word:
            cleaned_words.append(cleaned_word.strip())  # Optionally convert to lowercase
    return cleaned_words
def clean_single_word(word: str) -> str:
    """
    Cleans a single word by removing unwanted characters.

    Args:
        word (str): The word to be cleaned.

    Returns:
        str: The cleaned word.
    """
    # Remove unwanted characters using the translation table
    cleaned_word = word.translate(translation_table)
    return cleaned_word.strip()  # Optionally convert to lowercase if needed


async def extract_words_from_docx(file: UploadFile) -> List[str]:
    """
    Extracts words from a .docx file.

    Args:
        file (UploadFile): The uploaded .docx file.

    Returns:
        List[str]: A list of words extracted from the file.
    """
    try:
        # Check file extension (more reliable than content_type)
        if not file.filename.endswith('.docx'):
            logger.warning(f"Invalid DOCX file extension: {file.filename}")
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .docx file.")

        # Read file content into memory to handle SpooledTemporaryFile
        file_content = await file.read()
        
        # Create BytesIO object for DocxDocument to use
        bytes_io = BytesIO(file_content)
        bytes_io.seek(0)  # Reset to beginning
        
        # Open the document
        document = DocxDocument(bytes_io)
        words = []
        
        # Extract words from all paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                words.extend(paragraph.text.split())
        
        # Also extract from tables if any
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        words.extend(cell.text.split())
        
        logger.info(f"Successfully extracted {len(words)} words from DOCX file: {file.filename}")
        return words
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reading DOCX file {file.filename}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=400, 
            detail=f"Error processing DOCX file. File may be corrupted or invalid: {str(e)}"
        )


async def extract_words(file: UploadFile) -> List[str]:
    """
    Extracts words from either a .txt or .docx file.

    Args:
        file (UploadFile): The uploaded file.

    Returns:
        List[str]: A list of words extracted from the file.
    """
    try:
        # Validate file name and extension
        if not file.filename:
            raise HTTPException(status_code=400, detail="File name is missing.")
        
        filename_lower = file.filename.lower()
        
        if filename_lower.endswith('.txt'):
            return await extract_words_from_txt(file)
        elif filename_lower.endswith('.docx'):
            return await extract_words_from_docx(file)
        else:
            logger.warning(f"Unsupported file type: {file.filename}")
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a .txt or .docx file.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting words from file {file.filename}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


async def filter_words_from_file(file: UploadFile) -> List[str]:
    """
    Extracts words from the uploaded file and filters out the words already present in the Bloom filter.

    Args:
        file (UploadFile): The uploaded file.
        db (Session): The database session.
        bloom_filter (BloomWordFilter): An instance of the Bloom filter.

    Returns:
        List[str]: A list of words that are not present in the Bloom filter.
    """
    try:
        # Extract words based on file type
        words = await extract_words(file)
        
        if not words:
            logger.warning(f"No words extracted from file: {file.filename}")
            return []
        
        words = clean_words(words)
        
        if not words:
            logger.warning(f"No valid words after cleaning from file: {file.filename}")
            return []

        # Filter missing words using the Bloom filter
        missing_words = filter_missing_words(words=words)
        logger.info(f"Filtered {len(missing_words)} missing words from {len(words)} total words")

        return missing_words
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error filtering words from file {file.filename}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


async def filter_missing_words_from_list(words: List[str]) -> List[str]:
    """
    Filters a list of words to remove those that are already present in a Bloom filter.

    Args:
        words (List[str]): A list of words to be filtered.

    Returns:
        List[str]: A list of words that are not present in the Bloom filter.

    This function takes a list of words as input and uses the `filter_missing_words` function
    to filter out the words that are already present in a Bloom filter. The filtered words are
    then returned as a new list.
    """
    # Filter missing words using the Bloom filter
    missing_words = filter_missing_words(words=words)

    return missing_words


async def count_word_frequency(file: UploadFile) -> dict:
    """
    Counts the frequency of each word in the uploaded .txt or .docx file.

    Args:
        file (UploadFile): The uploaded file.

    Returns:
        dict: A dictionary where keys are words and values are their frequencies.
    """
    try:
        # Extract words based on file type
        words = await extract_words(file)
        
        if not words:
            logger.warning(f"No words extracted from file: {file.filename}")
            return {}
        
        words = clean_words(words)
        
        if not words:
            logger.warning(f"No valid words after cleaning from file: {file.filename}")
            return {}

        # Count the frequency of each word using Counter
        word_frequency = Counter(words)
        logger.info(f"Counted word frequency for {len(word_frequency)} unique words from {file.filename}")

        return dict(word_frequency)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error counting word frequency from file {file.filename}: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
