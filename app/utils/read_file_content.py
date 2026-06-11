from __future__ import annotations

import re
from collections import Counter
from typing import Any, Dict, List

import unicodedata
from docx import Document

from config.logger_config import setup_logger
from utils.kannada_word_clean import clean_kannada_word

logger = setup_logger(module_name=__name__)


class FileWordProcessor:
    """
    Process TXT and DOCX files and return
    word-frequency data after Kannada word cleaning.
    """

    # -----------------------------------------
    # PUBLIC API
    # -----------------------------------------

    @staticmethod
    def process_txt(file_path: str) -> List[Dict[str, Any]]:
        logger.info(f"TXT processing started | file={file_path}")

        text = FileWordProcessor._read_txt(file_path)
        result = FileWordProcessor._process_text(text)

        logger.info(
            f"TXT processing completed | "
            f"file={file_path} | "
            f"unique_words={len(result)}"
        )

        return result

    @staticmethod
    def process_docx(file_path: str) -> list[Dict[str, Any]]:
        logger.info(f"DOCX processing started | file={file_path}")

        text = FileWordProcessor._read_docx(file_path)
        result = FileWordProcessor._process_text(text)

        logger.info(
            f"DOCX processing completed | "
            f"file={file_path} | "
            f"unique_words={len(result)}"
        )

        return result

    # -----------------------------------------
    # FILE READERS
    # -----------------------------------------

    @staticmethod
    def _read_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()

            logger.info(
                f"TXT file loaded successfully | "
                f"file={file_path} | "
                f"characters={len(content)}"
            )

            return content

        except Exception:
            logger.exception(f"Failed to read TXT file | file={file_path}")
            raise

    @staticmethod
    def _read_docx(file_path: str) -> str:
        try:
            document = Document(file_path)

            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

            logger.info(
                f"DOCX file loaded successfully | "
                f"file={file_path} | "
                f"paragraphs={len(document.paragraphs)} | "
                f"characters={len(text)}"
            )

            return text

        except Exception:
            logger.exception(f"Failed to read DOCX file | file={file_path}")
            raise

    # -----------------------------------------
    # CORE PROCESSING
    # -----------------------------------------

    @staticmethod
    def _process_text(text: str) -> list[Dict[str, Any]]:
        if not text:
            logger.warning("Empty text received for processing")
            return []

        logger.debug(f"Starting text processing | characters={len(text)}")

        # Normalize Unicode
        text = unicodedata.normalize("NFC", text)

        # Extract tokens
        raw_words = re.findall(r"\S+", text)

        logger.info(f"Tokenization completed | raw_words={len(raw_words)}")

        cleaned_words: List[str] = []
        skipped_words = 0

        for word in raw_words:
            cleaned_word = clean_kannada_word(word)

            if cleaned_word:
                cleaned_words.append(cleaned_word)
            else:
                skipped_words += 1

        logger.info(
            f"Word cleaning completed | "
            f"valid_words={len(cleaned_words)} | "
            f"skipped_words={skipped_words}"
        )

        frequency_counter = Counter(cleaned_words)

        logger.info(
            f"Frequency analysis completed | " f"unique_words={len(frequency_counter)}"
        )

        result = [
            {
                "word": word,
                "frequency": frequency,
            }
            for word, frequency in frequency_counter.items()
        ]

        logger.info(
            f"Text processing completed | "
            f"total_words={len(cleaned_words)} | "
            f"unique_words={len(result)} | "
            f"skipped_words={skipped_words}"
        )

        return result
